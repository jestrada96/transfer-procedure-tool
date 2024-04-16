from docx import Document
from classes.node import Node
from excelData import pits
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

class DocWriter():
    def __init__(self, name="MyDoc"):
        self.name = name
        self.doc = Document()
        run = self.doc.add_heading("", 1).add_run()
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x42,0x42,0x42)
        run.add_text(self.name)
        
    def makeSection(self, name, instruction = None):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.space_before = Pt(1)
        section_name = paragraph.add_run(name)
        section_name.font.bold = True
        section_name.font.name = "Times New Roman"
        section_name.font.size = Pt(11)
        prompt = paragraph.add_run(instruction)
        prompt.font.name = "Times New Roman"
        prompt.font.size = Pt(11)
        return paragraph

    def save(self, filename= "PrintedRoute.docx"):
        self.doc.save(filename)

    def buildDocument(self, route):
        
        route_list = self.makeSection("Route List: ", "Valves in route (reference only):")
        used_pits = set()
        for node in route:
            used_pits.add(node.pit)
            if node.show:
                route_list.add_run("\n")
                route_list.add_run(node.EIN())
        
        heaterEINs = self.makeSection("Section 5.5.3 heaters: " ,"Replace existing data with the following:")
        for pit in used_pits:
            for heater in pits[pit].heaters:
                heaterEINs.add_run("\n")      
                heaterEINs.add_run(heater)
                heaterEINs.add_run("\t")
                heaterEINs.add_run("\t")
                heaterEINs.add_run(pits[pit].nacePMID)
        
        pits579 = self.makeSection("Steps 5.17.9: ","Replace existing data with the following:")
        checklist1 = self.makeSection("Checklist 1: ","Replace list with:")
        checklist3 = self.makeSection("Checklist 3:","")
        checklist4 = self.makeSection("Checklist 4:","")
        checklist5 = self.makeSection("Checklist 5:","")
        checklist6 = self.makeSection("Checklistc6:","")
        checklist7TF = self.makeSection("Checklist 7 - TFSPS Temperature Equipment Checks")
        for pit in used_pits:
            for tfsps , pmid in zip(pits[pit].tfsps,pits[pit].tfsps_pmid):
                checklist7TF.add_run("\n")
                checklist7TF.add_run(tfsps)
                checklist7TF.add_run("\t")
                checklist7TF.add_run("\t")
                checklist7TF.add_run(pmid)

        checklist7N = self.makeSection("Checklist 7 - NACE Inspection:")
        for pit in used_pits:
            checklist7N.add_run("\n")
            checklist7N.add_run(pits[pit].nace)
            checklist7N.add_run("\t")
            checklist7N.add_run("\t")
            checklist7N.add_run(pits[pit].nacePMID)

