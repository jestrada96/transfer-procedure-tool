from docx import Document
from indexed import IndexedOrderedDict
from docx.shared import Pt
from docx.shared import RGBColor
from procedure_data_tool.utils.valve2 import Valve2
from procedure_data_tool.utils.valve3 import Valve3
from procedure_data_tool.utils.line import Line

class DocWriter():
    def __init__(self, name="MyDoc"):
        self.name = name
        self.doc = Document()
        run = self.doc.add_heading("", 1).add_run()
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x42,0x42,0x42)
        run.add_text(self.name)
        
    def makeSection(self, name, instruction = None):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.space_before = Pt(1)
        section_name = paragraph.add_run(name)
        section_name.font.bold = True
        section_name.font.name = "Times New Roman"
        section_name.font.size = Pt(11)
        prompt = paragraph.add_run(instruction)
        prompt.font.name = "Times New Roman"
        prompt.font.size = Pt(11)
        return paragraph

    def save(self, filename= "PrintedRoute.docx"):
        self.doc.save(filename)

    def buildDocument(self, route, pits):
        used_pits = IndexedOrderedDict()
        used_jumpers = IndexedOrderedDict()
        used_lines = []
        for node in route:
            if node.pit and node.pit in pits:
                used_pits[node.pit] = pits[node.pit] 
                used_pits[node.pit].add_used_component(node)
            if node.onJumper:
                jumper = (node.pit, node.jumper)
                used_jumpers[jumper] = None

        # Add tank instead of TSR_Structure. Edit string??
        # Add rest of components in route.
        wlps_text = self.makeSection("Route Description: ", "use description for Waste Leak Path Screen")
        wlps_text.add_run("\n")
        sending_tank = used_pits.values()[0].tsr_structure[:-3] + "1" + used_pits.values()[0].tsr_structure[-3:-1]
        receiving_tank = used_pits.values()[-1].tsr_structure[:-3] + "1" + used_pits.values()[-1].tsr_structure[-3:-1]
        wlps_text.add_run(f"Waste from tank {sending_tank} will be transferred using {route[0]}, routed through ")
        for pit, line in zip(used_pits,used_lines):
            wlps_text.add_run(f"{pit} jumpers, ")
            wlps_text.add_run(f"{line.ein[-6:]}, ")    
        wlps_text.add_run(f" finally discharging into tank {receiving_tank}'s head space through the drop leg at {route[-1]}.")
        route_list = self.makeSection("Valves in Route (reference only): ", "DVI Credited YES/NO/POSition dependent")
        for node in route:
            if node.show:
                route_list.add_run("\n")
                route_list.add_run(node.EIN())
        heaterEINs = self.makeSection("Section 5.5.3 heaters: " ,"Replace existing data with the following:")
        for pit in used_pits.values():
            for heater in pit.in_pit_heaters:
                heaterEINs.add_run("\n")
                heaterEINs.add_run(heater)
                heaterEINs.add_run("\t \t")      
                heaterEINs.add_run(pit.pit_nace)
        pits5179 = self.makeSection("Steps 5.17.9: ","Replace existing data with the following:")
        for pit in used_pits.values():
            pits5179.add_run("\n")
            pits5179.add_run(pit.tsr_structure)
        checklist1 = self.makeSection("Checklist 1: ","Replace list with:")
        for jumper in used_jumpers:
            checklist1.add_run("\n")
            checklist1.add_run(jumper[0])
            checklist1.add_run("\t \t \t")
            checklist1.add_run("Jumper: ").font.bold = True
            checklist1.add_run(jumper[1])
        checklist3 = self.makeSection("*IN DEVELOPMENT* Checklist 3: Transfer Valving","")
        for pit in used_pits.values():
            checklist3.add_run("\n")
            # checklist3.add_run(pit.pit_nace[0:6]).bold = True
            checklist3.add_run(pit.pit_nace).bold = True
            checklist3.add_run(" Tank Farm").bold = True
            for component in pit.components:
                if (type(component) == Valve3 or type(component) == Valve2 ):
                    checklist3.add_run("\n")
                    checklist3.add_run(component.EIN())
                    checklist3.add_run("\t \t")
                    checklist3.add_run(component.position)
            checklist3.add_run("\n")
            checklist3.add_run("Confirm open route: ").bold = True
            checklist3.add_run(pit.pit_nace).bold = True
            checklist3.add_run("\n")
        checklist4 = self.makeSection("Checklist 4: Checklist 4 - Flush Transfer Route to Transfer Pump Valving","")
        checklist5 = self.makeSection("Checklist 5: Checklist 4 - Flush Transfer Route to Receiving Tank Valving","")
        checklist6 = self.makeSection("Checklist 6: Return to Transfer Valving","")
        checklist7LD = self.makeSection("Checklist 7 - Tank pit/Structure Leak Detection")
        checklist7TF = self.makeSection("Checklist 7 - TFSPS Temperature Equipment Checks")
        for pit in used_pits.values():
            for tfsps , pmid in zip(pit.tfsps_transmitters, pit.tfsps_pmids):
                checklist7TF.add_run("\n")
                checklist7TF.add_run(tfsps)
                checklist7TF.add_run("\t \t")
                checklist7TF.add_run(pmid)
        checklist7D = self.makeSection("Checklist 7 - Drain Seal Assemblies:")
        for pit in used_pits.values():
            checklist7D.add_run("\n")
            checklist7D.add_run(pit.drain_seal_location)
            checklist7D.add_run("\t \t")
            checklist7D.add_run(pit.drain_seal_position)
        checklist7N = self.makeSection("Checklist 7 - NACE Inspection:")
        for pit in used_pits.values():
            checklist7N.add_run("\n")
            checklist7N.add_run(pit.pit_nace)
            checklist7N.add_run("\t \t")
            checklist7N.add_run(pit.pit_nace_pmid)



